#!/usr/bin/env python3
import docx
from docx.enum.text import WD_UNDERLINE
from docx.oxml.ns import qn
from docx.shared import Pt
import sys, os, re

fname = sys.argv[1]
tname = os.path.basename(fname.replace(".tsv", ".docx"))

doc = docx.Document()
doc.settings.element.append(
    docx.oxml.OxmlElement('w:hideSpellingErrors'))
doc.settings.element.append(
    docx.oxml.OxmlElement('w:hideGrammaticalErrors'))
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(16)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
for line in open(fname, encoding="utf-8"):
    line = line.rstrip()
    if line.startswith("#") or re.match(r"^[^\d⓪①-⑨ⓐⓑ]+$", line):
        h = doc.add_heading("", level=0)
        run = h.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(20)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    else:
        p = doc.add_paragraph()
        sub = False
        for i in line:
            if i == "{":
                sub = True
            elif i == "}":
                sub = False
            elif i == "-":
                r.font.underline = WD_UNDERLINE.SINGLE
            elif i == "=":
                r.font.underline = WD_UNDERLINE.DOUBLE
            elif sub:
                r = p.add_run(i)
                r.font.subscript = True
            else:
                r = p.add_run(i)
                if i == "□":
                    r.font.name = "宋体"
doc.save(tname)